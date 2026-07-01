"""
chronos.report
==============
Builds the case-study Word document (``case.docx``) that collates *all*
generated engineering output: input/site data, model set-up, mesh/discretisation,
solution controls, every diagnostic figure, every results table, validation and
conclusions.

Uses ``python-docx``. Figures are embedded from disk; tables are rendered from
pandas DataFrames. No colour styling relies on black.
"""

from __future__ import annotations
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INK = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x1B, 0x6C, 0xA8)


class CaseReport:
    def __init__(self, title, subtitle="", author="", affiliation=""):
        self.doc = Document()
        self._base_style()
        self._title_block(title, subtitle, author, affiliation)

    def _base_style(self):
        st = self.doc.styles["Normal"]
        st.font.name = "Calibri"
        st.font.size = Pt(10.5)
        st.font.color.rgb = INK

    def _title_block(self, title, subtitle, author="", affiliation=""):
        h = self.doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run(title); r.bold = True; r.font.size = Pt(19)
        r.font.color.rgb = INK
        if subtitle:
            s = self.doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rs = s.add_run(subtitle); rs.italic = True; rs.font.size = Pt(11.5)
            rs.font.color.rgb = ACCENT
        if author:
            a = self.doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ra = a.add_run(author); ra.bold = True; ra.font.size = Pt(12.5)
            ra.font.color.rgb = INK
        if affiliation:
            af = self.doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER
            raf = af.add_run(affiliation); raf.italic = True; raf.font.size = Pt(10)
            raf.font.color.rgb = ACCENT
        self.doc.add_paragraph()

    # -- structural helpers ------------------------------------------------ #
    def h1(self, text):
        p = self.doc.add_heading(level=1)
        r = p.add_run(text); r.font.color.rgb = ACCENT; r.bold = True
        return p

    def h2(self, text):
        p = self.doc.add_heading(level=2)
        r = p.add_run(text); r.font.color.rgb = INK; r.bold = True
        return p

    def para(self, text, italic=False, bold=False):
        p = self.doc.add_paragraph()
        r = p.add_run(text); r.italic = italic; r.bold = bold
        r.font.color.rgb = INK
        return p

    def bullets(self, items):
        for it in items:
            p = self.doc.add_paragraph(style="List Bullet")
            r = p.add_run(it); r.font.color.rgb = INK

    def numbered(self, items):
        for it in items:
            p = self.doc.add_paragraph(style="List Number")
            r = p.add_run(it); r.font.color.rgb = INK

    def figure(self, path, caption, width=6.2):
        if not os.path.exists(path):
            self.para(f"[missing figure: {path}]", italic=True); return
        self.doc.add_picture(path, width=Inches(width))
        last = self.doc.paragraphs[-1]; last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
        r.font.color.rgb = ACCENT

    def table(self, df: pd.DataFrame, caption=None, index=False, float_fmt="{:.4g}"):
        if caption:
            cp = self.doc.add_paragraph()
            r = cp.add_run(caption); r.bold = True; r.font.size = Pt(9.5)
            r.font.color.rgb = INK
        cols = list(df.columns)
        if index:
            cols = [df.index.name or ""] + cols
        t = self.doc.add_table(rows=1, cols=len(cols))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for j, c in enumerate(cols):
            run = hdr[j].paragraphs[0].add_run(str(c))
            run.bold = True; run.font.size = Pt(8.5); run.font.color.rgb = INK
        for idx, row in df.iterrows():
            cells = t.add_row().cells
            off = 0
            if index:
                cells[0].paragraphs[0].add_run(str(idx)).font.size = Pt(8.5)
                off = 1
            for j, c in enumerate(df.columns):
                v = row[c]
                if isinstance(v, float):
                    txt = float_fmt.format(v)
                else:
                    txt = str(v)
                run = cells[j+off].paragraphs[0].add_run(txt)
                run.font.size = Pt(8.5); run.font.color.rgb = INK
        self.doc.add_paragraph()

    def equation(self, path, width_in=6.3, number=None):
        """Embed a LaTeX-rendered equation image at its native (standard) size."""
        if not os.path.exists(path):
            self.para(f"[missing equation image: {path}]", italic=True); return
        self.doc.add_picture(path, width=Inches(width_in))
        last = self.doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if number:
            cap = self.doc.add_paragraph()
            r = cap.add_run(f"({number})"); r.italic = True; r.font.size = Pt(8.5)
            r.font.color.rgb = ACCENT

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)
        return path
